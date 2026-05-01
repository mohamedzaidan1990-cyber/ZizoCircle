from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as OxmlEl
import copy

# ── Colour constants ──────────────────────────────────────────────────────────
PURPLE      = RGBColor(0x7B, 0x5C, 0xF6)
DARK_NAVY   = RGBColor(0x0D, 0x0B, 0x1E)
LIGHT_PURPLE= RGBColor(0xA7, 0x8B, 0xFA)
BODY_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
MUTED_GREY  = RGBColor(0x6B, 0x72, 0x80)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIME        = RGBColor(0x16, 0xA3, 0x4A)

OUTPUT_PATH = r"C:\Users\User\Downloads\ZizoCircle_Snoonu_Application.docx"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_run(run, text, font_name="Arial", size_pt=9, bold=False,
            italic=False, color=None):
    run.text = text
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=0, left_indent=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment = align
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    fmt.line_spacing = None
    if left_indent is not None:
        fmt.left_indent = Cm(left_indent)
    return p

def add_hr(doc, color="7B5CF6", size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def clear_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'),   'none')
        border.set(qn('w:sz'),    '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def clear_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'),   'none')
        border.set(qn('w:sz'),    '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_width(cell, cm_val):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing tcW
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    # Convert cm to twentieths of a point (twips): 1 cm = 567 twips
    twips = int(cm_val * 567)
    tcW.set(qn('w:w'),    str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)

def section_heading(doc, text, space_before=6):
    p = add_para(doc, space_before=space_before, space_after=2)
    r = p.add_run(text)
    set_run(r, text, size_pt=10, bold=True, color=PURPLE)
    return p

def body_para(doc, text, space_after=6, left_indent=None, line_spacing=1.15):
    p = add_para(doc, space_after=space_after, left_indent=left_indent)
    r = p.add_run(text)
    set_run(r, text, size_pt=9.5, color=BODY_DARK)
    if line_spacing:
        from docx.shared import Pt as _Pt
        p.paragraph_format.line_spacing = _Pt(9.5 * line_spacing)
    return p

def bullet_para(doc, text, space_after=1, left_indent=0.5):
    p = add_para(doc, space_after=space_after, left_indent=left_indent)
    r = p.add_run(text)
    set_run(r, text, size_pt=9.5, color=BODY_DARK)
    return p

# ── Build document ────────────────────────────────────────────────────────────
doc = Document()

# Page setup
section = doc.sections[0]
from docx.shared import Cm as _Cm
section.page_width  = _Cm(21)
section.page_height = _Cm(29.7)
section.top_margin    = _Cm(1.5)
section.bottom_margin = _Cm(1.5)
section.left_margin   = _Cm(2.0)
section.right_margin  = _Cm(2.0)

# Remove default paragraph spacing from Normal style
from docx.styles.style import _ParagraphStyle
normal_style = doc.styles['Normal']
normal_style.paragraph_format.space_before = Pt(0)
normal_style.paragraph_format.space_after  = Pt(0)

# ── 1. Letterhead block ───────────────────────────────────────────────────────
# Usable width: 21 - 2 - 2 = 17 cm
header_table = doc.add_table(rows=1, cols=2)
header_table.style = 'Table Grid'
clear_table_borders(header_table)
header_table.alignment = WD_TABLE_ALIGNMENT.LEFT

left_cell  = header_table.cell(0, 0)
right_cell = header_table.cell(0, 1)

# Set widths: 70% = 11.9 cm, 30% = 5.1 cm
set_cell_width(left_cell,  11.9)
set_cell_width(right_cell,  5.1)

clear_cell_borders(left_cell)
clear_cell_borders(right_cell)

# Remove default empty paragraph from cells then build content
for cell in [left_cell, right_cell]:
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

# Left cell content
def add_cell_para(cell, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=0, left_indent=None):
    p = cell.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment    = align
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if left_indent is not None:
        fmt.left_indent = Cm(left_indent)
    return p

# Line 1: Zizo Circle
p = add_cell_para(left_cell, space_after=1)
r = p.add_run("Zizo Circle")
set_run(r, "Zizo Circle", font_name="Arial Black", size_pt=28, bold=True, color=PURPLE)

# Line 2: by Bayline Holding
p = add_cell_para(left_cell, space_after=1)
r = p.add_run("by Bayline Holding W.L.L.")
set_run(r, "by Bayline Holding W.L.L.", size_pt=10, color=MUTED_GREY)

# Line 3: blank 4pt
p = add_cell_para(left_cell, space_after=4)
p.add_run("")

# Line 4: CEO line
p = add_cell_para(left_cell, space_after=1)
r = p.add_run("CEO: Mohamed Zaidan  ·  Lusail, Doha, Qatar")
set_run(r, r.text, size_pt=9, color=BODY_DARK)

# Line 5: contact
p = add_cell_para(left_cell, space_after=0)
r = p.add_run("+974 7727 7292  ·  thediamondexpert974@gmail.com")
set_run(r, r.text, size_pt=9, color=MUTED_GREY)

# Right cell content
# Line 1: COMPANY OVERVIEW
p = add_cell_para(right_cell, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=1)
r = p.add_run("COMPANY OVERVIEW")
set_run(r, r.text, size_pt=8, bold=True, color=PURPLE)

# Line 2: Snoonu Incubator Application
p = add_cell_para(right_cell, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=1)
r = p.add_run("Snoonu Incubator Application")
set_run(r, r.text, size_pt=8, color=MUTED_GREY)

# Line 3: May 2026
p = add_cell_para(right_cell, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
r = p.add_run("May 2026")
set_run(r, r.text, size_pt=8, color=MUTED_GREY)

# ── 2. Purple divider ─────────────────────────────────────────────────────────
add_hr(doc, color="7B5CF6", size=12)

# ── 3. MVP LIVE badge ─────────────────────────────────────────────────────────
p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=8)
badge_text = "● MVP IS LIVE  ·  Android APK Available  ·  Real Users Registered  ·  Firebase Backend Active"
r = p.add_run(badge_text)
set_run(r, badge_text, size_pt=9, bold=True, color=LIME)

# ── 4. WHAT WE DO ─────────────────────────────────────────────────────────────
section_heading(doc, "WHAT WE DO", space_before=6)
body_para(doc,
    "Zizo Circle is Qatar’s first AI-powered social interest matching app. "
    "Users build a profile, choose from 36 interest categories, and Zizo’s AI "
    "matches them with compatible people nearby. When matched users visit partner "
    "venues together, they earn 2× cashback rewards automatically — turning "
    "social connections into real financial benefits.",
    space_after=6)

# ── 5. CURRENT STATUS ─────────────────────────────────────────────────────────
section_heading(doc, "CURRENT STATUS", space_before=6)
bullets_status = [
    "✅  Android APK live and available for download via Expo",
    "✅  Real users registered — Firebase Auth + Firestore backend active",
    "✅  Zizo AI companion powered by Claude (Anthropic)",
    "✅  36 interest categories · AI matching engine · Cashback wallet",
    "✅  React Native + Expo codebase — iOS build ready post-funding",
]
for b in bullets_status:
    bullet_para(doc, b, space_after=1, left_indent=0.5)

# ── 6. BUSINESS MODEL ─────────────────────────────────────────────────────────
section_heading(doc, "BUSINESS MODEL", space_before=6)

biz_table = doc.add_table(rows=1, cols=3)
biz_table.style = 'Table Grid'
clear_table_borders(biz_table)
biz_table.alignment = WD_TABLE_ALIGNMENT.LEFT

col_widths = [5.67, 5.67, 5.66]
biz_data = [
    ("10%",       ["Venue commission on", "cashback transactions"]),
    ("QAR 50/mo", ["Premium user",        "subscriptions"]),
    ("Featured",  ["Venue listing &",     "placement fees"]),
]

for i, (big_text, small_lines) in enumerate(biz_data):
    cell = biz_table.cell(0, i)
    set_cell_width(cell, col_widths[i])
    clear_cell_borders(cell)
    # Remove default paragraph
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    # Big number
    p = add_cell_para(cell, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    r = p.add_run(big_text)
    set_run(r, big_text, size_pt=18, bold=True, color=PURPLE)

    # Small lines
    for line in small_lines:
        p = add_cell_para(cell, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        r = p.add_run(line)
        set_run(r, line, size_pt=8, color=MUTED_GREY)

# ── 7. MARKET OPPORTUNITY ─────────────────────────────────────────────────────
section_heading(doc, "MARKET OPPORTUNITY", space_before=8)
body_para(doc,
    "Qatar: 1.5M+ socially active adults — underserved by existing social platforms. "
    "Year 1 target: 100,000 users · QAR 7.2M revenue. "
    "Year 2: GCC expansion beginning with UAE (3.5M+ target market). "
    "The Qatar National Vision 2030 prioritises community engagement and digital "
    "innovation — Zizo Circle is built for this moment.",
    space_after=6)

# ── 8. OUR ASK FROM SNOONU ────────────────────────────────────────────────────
section_heading(doc, "OUR ASK FROM SNOONU", space_before=8)

ask_table = doc.add_table(rows=1, cols=2)
ask_table.style = 'Table Grid'
clear_table_borders(ask_table)
ask_table.alignment = WD_TABLE_ALIGNMENT.LEFT

# 55% = 9.35 cm, 45% = 7.65 cm
left_ask  = ask_table.cell(0, 0)
right_ask = ask_table.cell(0, 1)
set_cell_width(left_ask,  9.35)
set_cell_width(right_ask, 7.65)
clear_cell_borders(left_ask)
clear_cell_borders(right_ask)

for cell in [left_ask, right_ask]:
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

# Left column
p = add_cell_para(left_ask, space_after=1)
r = p.add_run("QAR 1,255,000")
set_run(r, r.text, size_pt=22, bold=True, color=PURPLE)

p = add_cell_para(left_ask, space_after=1)
r = p.add_run("Seed funding for 23.9% equity")
set_run(r, r.text, size_pt=9, color=MUTED_GREY)

p = add_cell_para(left_ask, space_after=4)
p.add_run("")

left_bullets = [
    "·  Full iOS + Android public launch",
    "·  Marketing & influencer campaigns",
    "·  Venue partnership scaling to 100+",
]
for b in left_bullets:
    p = add_cell_para(left_ask, space_after=1)
    r = p.add_run(b)
    set_run(r, b, size_pt=9, color=BODY_DARK)

# Right column
p = add_cell_para(right_ask, space_after=1)
r = p.add_run("Beyond Capital")
set_run(r, r.text, size_pt=10, bold=True, color=PURPLE)

right_bullets = [
    "·  Snoonu mentor & network access",
    "·  Ecosystem integration opportunity",
    "·  Co-marketing & brand exposure",
]
for b in right_bullets:
    p = add_cell_para(right_ask, space_after=1)
    r = p.add_run(b)
    set_run(r, b, size_pt=9, color=BODY_DARK)

p = add_cell_para(right_ask, space_after=0)
r = p.add_run("Could Zizo Circle complement Snoonu’s lifestyle platform?")
set_run(r, r.text, size_pt=8.5, italic=True, color=MUTED_GREY)

# ── 9. Second divider ─────────────────────────────────────────────────────────
add_hr(doc, color="7B5CF6", size=12)

# ── 10. Footer ────────────────────────────────────────────────────────────────
p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=0)
footer_text = ("Zizo Circle · Bayline Holding W.L.L. · Lusail, Doha, Qatar "
               "· thediamondexpert974@gmail.com · +974 7727 7292")
r = p.add_run(footer_text)
set_run(r, footer_text, size_pt=8, color=MUTED_GREY)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)

import os
size = os.path.getsize(OUTPUT_PATH)
print(f"Document saved: {OUTPUT_PATH}")
print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")
